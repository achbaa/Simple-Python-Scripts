import os
import sys
import docx
import asyncio
import aiohttp
from openai import AsyncOpenAI
import argparse
from pathlib import Path
from dotenv import load_dotenv
from datetime import datetime
from tqdm import tqdm
import tiktoken

# Load environment variables from .env file
load_dotenv()

# Get API key from environment variable
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    print("Error: OPENAI_API_KEY not found in environment variables.")
    print("Please create a .env file with your OpenAI API key or set it as an environment variable.")
    sys.exit(1)

# Initialize OpenAI client
client = AsyncOpenAI(api_key=api_key)

def read_text_file(file_path):
    """Read content from a text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            # Check if content is empty or just whitespace
            if not content.strip():
                print(f"Warning: {file_path} appears to be empty or contains only whitespace.")
            return content
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return f"Error reading file: {e}"

def read_docx_file(file_path):
    """Read content from a Word document."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        content = '\n'.join(full_text)
        # Check if content is empty or just whitespace
        if not content.strip():
            print(f"Warning: {file_path} appears to be empty or contains only whitespace.")
        return content
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return f"Error reading file: {e}"

def read_file(file_path):
    """Read content from a file based on its extension."""
    if file_path.lower().endswith('.txt'):
        return read_text_file(file_path)
    elif file_path.lower().endswith(('.docx', '.doc')):
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_path}")

def count_tokens(text, model="gpt-4"):
    """Count the number of tokens in the text."""
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except KeyError:
        # Fallback to cl100k_base encoding for newer models not yet in tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))

def chunk_text(text, max_tokens=3000, model="gpt-4"):
    """Split text into chunks that don't exceed max_tokens."""
    try:
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        # Fallback to cl100k_base encoding for newer models not yet in tiktoken
        encoding = tiktoken.get_encoding("cl100k_base")
        
    tokens = encoding.encode(text)
    
    chunks = []
    current_chunk = []
    current_count = 0
    
    for token in tokens:
        if current_count + 1 > max_tokens:
            chunks.append(encoding.decode(current_chunk))
            current_chunk = [token]
            current_count = 1
        else:
            current_chunk.append(token)
            current_count += 1
            
    if current_chunk:
        chunks.append(encoding.decode(current_chunk))
        
    return chunks

async def analyze_transcript(transcript, file_name, model="gpt-4"):
    """Analyze a transcript using OpenAI API."""
    # Check if transcript is empty or contains error message
    if not transcript.strip() or transcript.startswith("Error reading file:"):
        return f"# Error Analyzing {file_name}\n\nThe transcript appears to be empty or could not be read properly."
    
    # Check if transcript is too long and chunk if necessary
    token_count = count_tokens(transcript, model)
    
    if token_count > 3000:
        print(f"  Transcript is long ({token_count} tokens). Processing in chunks...")
        chunks = chunk_text(transcript, 3000, model)
        chunk_analyses = []
        
        # Process chunks concurrently with rate limiting
        semaphore = asyncio.Semaphore(4)  # Limit concurrent API calls
        
        async def process_chunk(chunk, chunk_index):
            async with semaphore:
                chunk_prompt = f"""
                You are analyzing part {chunk_index+1} of {len(chunks)} of an expert interview transcript. The file name is {file_name}.
                For this chunk, extract:
                
                1. Key points
                2. Notable quotes (with proper attribution)
                3. Emerging themes
                
                Format your response in markdown.
                
                Here is the transcript chunk:
                {chunk}
                """
                
                try:
                    response = await client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are an expert at analyzing interview transcripts and extracting key information."},
                            {"role": "user", "content": chunk_prompt}
                        ],
                        max_tokens=1000
                    )
                    return response.choices[0].message.content
                except Exception as e:
                    print(f"  Error analyzing chunk {chunk_index+1} of {file_name}: {e}")
                    return f"Error analyzing chunk {chunk_index+1}: {e}"
        
        # Process all chunks concurrently
        tasks = [process_chunk(chunk, i) for i, chunk in enumerate(chunks)]
        chunk_analyses = await asyncio.gather(*tasks)
        
        # Now synthesize the chunk analyses
        synthesis_prompt = f"""
        You are synthesizing analyses of multiple chunks from the same interview transcript. The file name is {file_name}.
        Based on these chunk analyses, provide:
        
        1. A concise summary (150-200 words)
        2. 5-7 key quotes (with proper attribution)
        3. 1-3 key themes or insights
        
        Format your response in markdown with clear headings for each section.
        
        Here are the chunk analyses:
        {"".join(chunk_analyses)}
        """
        
        try:
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an expert at synthesizing information from interview analyses."},
                    {"role": "user", "content": synthesis_prompt}
                ],
                max_tokens=1000
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"  Error synthesizing chunks for {file_name}: {e}")
            return f"# Error Analyzing {file_name}\n\nAn error occurred during synthesis: {e}"
    
    else:
        # Process as a single chunk
        prompt = f"""
        You are analyzing an expert interview transcript. The file name is {file_name}.
        Please analyze this transcript and provide:
        
        1. A concise summary (150-200 words)
        2. 5-7 key quotes (with proper attribution)
        3. 1-3 key themes or insights
        
        Format your response in markdown with clear headings for each section.
        
        Here is the transcript:
        {transcript}
        """
        
        try:
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an expert at analyzing interview transcripts and extracting key information."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"  Error analyzing transcript {file_name}: {e}")
            return f"# Error Analyzing {file_name}\n\nAn error occurred during analysis: {e}"

async def create_overall_synthesis(analyses, model="gpt-4"):
    """Create an overall synthesis of all interview analyses."""
    synthesis_prompt = f"""
    You are synthesizing analyses from multiple expert interviews.
    Based on these {len(analyses)} interview analyses, provide:
    
    1. An executive summary (200-300 words)
    2. Key themes across all interviews
    3. Notable insights and patterns
    4. Important contrasts or differences in perspectives
    
    Format your response in markdown with clear headings for each section.
    
    Here are the individual analyses:
    {"".join(analyses)}
    """
    
    try:
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert at synthesizing information from multiple interview analyses."},
                {"role": "user", "content": synthesis_prompt}
            ],
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error creating overall synthesis: {e}")
        return f"# Error Creating Overall Synthesis\n\nAn error occurred: {e}"

async def process_transcripts(transcript_files, model):
    """Process multiple transcripts concurrently."""
    semaphore = asyncio.Semaphore(4)  # Limit concurrent processing
    
    async def process_file(file_path):
        async with semaphore:
            file_name = os.path.basename(file_path)
            print(f"\nAnalyzing {file_name}...")
            
            try:
                transcript = read_file(file_path)
                preview = transcript[:100].replace('\n', ' ').strip()
                print(f"  Preview: {preview}...")
                
                if not transcript.strip():
                    print(f"  Warning: {file_name} appears to be empty.")
                    return f"# Error Analyzing {file_name}\n\nThe transcript appears to be empty."
                
                analysis = await analyze_transcript(transcript, file_name, model)
                print(f"Completed analysis of {file_name}")
                return analysis
            except Exception as e:
                print(f"Error processing {file_name}: {e}")
                return f"# Error Processing {file_name}\n\nAn error occurred: {e}"
    
    analyses = await asyncio.gather(*[process_file(file_path) for file_path in transcript_files])
    return analyses

async def async_main():
    """Async main function to run the script."""
    # Check if command-line arguments are provided
    if len(sys.argv) > 1:
        # Use argparse for command-line arguments
        parser = argparse.ArgumentParser(description='Analyze interview transcripts and generate summaries.')
        parser.add_argument('folder_path', help='Path to the folder containing interview transcripts')
        parser.add_argument('--output', help='Path to save the output file', default=None)
        parser.add_argument('--model', help='OpenAI model to use', default="gpt-4")
        parser.add_argument('--format', help='Output format (md or docx)', default="md")
        
        args = parser.parse_args()
        folder_path = args.folder_path
        output_path = args.output
        model = args.model
        output_format = args.format.lower()
    else:
        # Get input interactively
        user_input = get_user_input()
        folder_path = user_input["folder_path"]
        output_path = user_input["output_path"]
        model = user_input["model"]
        output_format = user_input["output_format"]
    
    if output_format not in ["md", "docx"]:
        print("Error: Output format must be 'md' or 'docx'")
        sys.exit(1)
    
    # Find all transcript files
    transcript_files = []
    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        if os.path.isfile(file_path) and file.lower().endswith(('.txt', '.docx', '.doc')):
            transcript_files.append(file_path)
    
    if not transcript_files:
        print(f"No transcript files found in {folder_path}. Please ensure files have .txt, .docx, or .doc extensions.")
        sys.exit(1)
    
    # Sort transcript files to ensure consistent processing order
    transcript_files.sort()
    
    print(f"\nFound {len(transcript_files)} transcript files. Beginning analysis...")
    print("Files to be processed:")
    for i, file_path in enumerate(transcript_files, 1):
        print(f"  {i}. {os.path.basename(file_path)}")
    
    # Analyze transcripts concurrently
    individual_analyses = await process_transcripts(transcript_files, model)
    all_analyses = [f"# Analysis of {os.path.basename(f)}\n\n{a}\n\n---\n\n" 
                   for f, a in zip(transcript_files, individual_analyses)]
    
    # Create overall synthesis
    print("\nCreating overall synthesis...")
    overall_synthesis = await create_overall_synthesis(individual_analyses, model)
    
    # Combine everything into final document
    final_document = f"""# Interview Transcript Synthesis
Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{overall_synthesis}

---

# Individual Interview Analyses

"""
    
    final_document += "\n\n".join(all_analyses)
    
    # Determine output file path
    if output_path:
        final_output_path = output_path
    else:
        # Create output directory if it doesn't exist
        script_dir = Path(__file__).parent
        output_dir = script_dir / "Output files"
        output_dir.mkdir(exist_ok=True)
        
        if output_format == "md":
            final_output_path = output_dir / "interview_synthesis.md"
        else:
            final_output_path = output_dir / "interview_synthesis.docx"
    
    # Write to file
    if output_format == "md":
        with open(final_output_path, 'w', encoding='utf-8') as f:
            f.write(final_document)
    else:
        # Create a Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading('Interview Transcript Synthesis', 0)
        doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph("---")
        
        # Parse markdown and add to document
        # This is a simple implementation - a more robust solution would use a markdown parser
        sections = overall_synthesis.split('\n#')
        for i, section in enumerate(sections):
            if i == 0:  # First section doesn't start with #
                lines = section.split('\n')
            else:
                lines = ('#' + section).split('\n')
                
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('> '):
                    p = doc.add_paragraph(line[2:])
                    p.style = 'Quote'
                elif line.strip() == '---':
                    doc.add_paragraph('---')
                elif line.strip():
                    doc.add_paragraph(line)
        
        doc.add_heading('Individual Interview Analyses', 1)
        
        # Add individual analyses
        for analysis in all_analyses:
            sections = analysis.split('\n#')
            for i, section in enumerate(sections):
                if i == 0:  # First section doesn't start with #
                    lines = section.split('\n')
                else:
                    lines = ('#' + section).split('\n')
                    
                for line in lines:
                    if line.startswith('# '):
                        doc.add_heading(line[2:], 2)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], 3)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], 4)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('> '):
                        p = doc.add_paragraph(line[2:])
                        p.style = 'Quote'
                    elif line.strip() == '---':
                        doc.add_paragraph('---')
                    elif line.strip():
                        doc.add_paragraph(line)
        
        doc.save(final_output_path)
    
    print(f"\nAnalysis complete! Results saved to {final_output_path}")

def get_user_input():
    """Get input from the user interactively."""
    print("\n=== Interview Transcript Synthesizer ===\n")
    
    # Get folder path
    while True:
        folder_path = input("Enter the path to the folder containing interview transcripts: ").strip()
        if os.path.isdir(folder_path):
            break
        else:
            print(f"Error: '{folder_path}' is not a valid directory. Please try again.")
    
    # Get output file path (optional)
    output_path = input("Enter the path for the output file (leave blank for default): ").strip()
    
    # Get model - expanded options
    available_models = ["gpt-4", "gpt-4o", "gpt-4o-mini", "gpt-3.5-turbo"]
    print("\nAvailable models:")
    for i, model in enumerate(available_models, 1):
        print(f"{i}. {model}")
    
    while True:
        model_choice = input(f"Select a model (1-{len(available_models)}, default is 1): ").strip()
        if not model_choice:
            model = available_models[0]
            break
        try:
            model_index = int(model_choice) - 1
            if 0 <= model_index < len(available_models):
                model = available_models[model_index]
                break
            else:
                print(f"Please enter a number between 1 and {len(available_models)}.")
        except ValueError:
            print("Please enter a valid number.")
    
    # Get output format
    print("\nAvailable output formats:")
    print("1. Markdown (.md)")
    print("2. Word Document (.docx)")
    
    while True:
        format_choice = input("Select an output format (1-2, default is 1): ").strip()
        if not format_choice:
            output_format = "md"
            break
        try:
            format_index = int(format_choice)
            if format_index == 1:
                output_format = "md"
                break
            elif format_index == 2:
                output_format = "docx"
                break
            else:
                print("Please enter either 1 or 2.")
        except ValueError:
            print("Please enter a valid number.")
    
    return {
        "folder_path": folder_path,
        "output_path": output_path,
        "model": model,
        "output_format": output_format
    }

if __name__ == "__main__":
    asyncio.run(async_main())
